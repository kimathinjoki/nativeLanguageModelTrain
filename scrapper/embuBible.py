import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import os
import re
from datetime import datetime

# Bible book codes with their chapter counts
BIBLE_BOOKS = [
    ("GEN", "Genesis", 50),      # Genesis (Kîam)
    ("EXO", "Exodus", 40),       # Exodus
    ("LEV", "Leviticus", 27),    # Leviticus (Ala)
    ("NUM", "Numbers", 36),      # Numbers
    ("DEU", "Deuteronomy", 34),  # Deuteronomy
    ("JOS", "Joshua", 24),       # Joshua
    ("JDG", "Judges", 21),       # Judges
    ("RUT", "Ruth", 4),          # Ruth
    ("1SA", "1 Samuel", 31),     # 1 Samuel
    ("2SA", "2 Samuel", 24),     # 2 Samuel
    ("1KI", "1 Kings", 22),      # 1 Kings
    ("2KI", "2 Kings", 25),      # 2 Kings
    ("1CH", "1 Chronicles", 29), # 1 Chronicles
    ("2CH", "2 Chronicles", 36), # 2 Chronicles
    ("EZR", "Ezra", 10),         # Ezra
    ("NEH", "Nehemiah", 13),     # Nehemiah
    ("EST", "Esther", 10),       # Esther
    ("JOB", "Job", 42),          # Job
    ("PSA", "Psalms", 150),      # Psalms
    ("PRO", "Proverbs", 31),     # Proverbs
    ("ECC", "Ecclesiastes", 12), # Ecclesiastes
    ("SNG", "Song of Songs", 8), # Song of Songs
    ("ISA", "Isaiah", 66),       # Isaiah
    ("JER", "Jeremiah", 52),     # Jeremiah
    ("LAM", "Lamentations", 5),  # Lamentations
    ("EZK", "Ezekiel", 48),      # Ezekiel
    ("DAN", "Daniel", 12),       # Daniel
    ("HOS", "Hosea", 14),        # Hosea
    ("JOL", "Joel", 3),          # Joel
    ("AMO", "Amos", 9),          # Amos
    ("OBA", "Obadiah", 1),       # Obadiah
    ("JON", "Jonah", 4),         # Jonah
    ("MIC", "Micah", 7),         # Micah
    ("NAM", "Nahum", 3),         # Nahum
    ("HAB", "Habakkuk", 3),      # Habakkuk
    ("ZEP", "Zephaniah", 3),     # Zephaniah
    ("HAG", "Haggai", 2),        # Haggai
    ("ZEC", "Zechariah", 14),    # Zechariah
    ("MAL", "Malachi", 4),       # Malachi
    ("MAT", "Matthew", 28),      # Matthew
    ("MRK", "Mark", 16),         # Mark
    ("LUK", "Luke", 24),         # Luke
    ("JHN", "John", 21),         # John
    ("ACT", "Acts", 28),         # Acts
    ("ROM", "Romans", 16),       # Romans
    ("1CO", "1 Corinthians", 16),# 1 Corinthians
    ("2CO", "2 Corinthians", 13),# 2 Corinthians
    ("GAL", "Galatians", 6),     # Galatians
    ("EPH", "Ephesians", 6),     # Ephesians
    ("PHP", "Philippians", 4),   # Philippians
    ("COL", "Colossians", 4),    # Colossians
    ("1TH", "1 Thessalonians", 5),# 1 Thessalonians
    ("2TH", "2 Thessalonians", 3),# 2 Thessalonians
    ("1TI", "1 Timothy", 6),     # 1 Timothy
    ("2TI", "2 Timothy", 4),     # 2 Timothy
    ("TIT", "Titus", 3),         # Titus
    ("PHM", "Philemon", 1),      # Philemon
    ("HEB", "Hebrews", 13),      # Hebrews
    ("JAS", "James", 5),         # James
    ("1PE", "1 Peter", 5),       # 1 Peter
    ("2PE", "2 Peter", 3),       # 2 Peter
    ("1JN", "1 John", 5),        # 1 John
    ("2JN", "2 John", 1),        # 2 John
    ("3JN", "3 John", 1),        # 3 John
    ("JUD", "Jude", 1),          # Jude
    ("REV", "Revelation", 22)    # Revelation
]

# Dictionary for quick lookup of chapter counts
CHAPTER_COUNTS = {code: count for code, _, count in BIBLE_BOOKS}

def scrape_bible_text(url):
    """
    Scrape Bible text from bible.com with improved pattern recognition
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Connection": "keep-alive"
    }
    
    print(f"Scraping: {url}")
    response = requests.get(url, headers=headers)
    
    if response.status_code != 200:
        print(f"Failed to retrieve the page: Status code {response.status_code}")
        return None
    
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Get chapter title from h1 (like "Kîam 1")
    chapter_title = None
    h1_elem = soup.find('h1')
    if h1_elem:
        chapter_title = h1_elem.text.strip()
        print(f"Found chapter title: {chapter_title}")
    
    # Extract book name and chapter number
    book_name = ""
    chapter_num = ""
    if chapter_title:
        parts = chapter_title.split()
        if len(parts) >= 2 and parts[-1].isdigit():
            book_name = " ".join(parts[:-1])
            chapter_num = parts[-1]
        else:
            book_name = chapter_title
    
    # Extract book code and chapter from URL
    url_match = re.search(r'/bible/\d+/([A-Z0-9]+)\.(\d+)', url)
    if url_match:
        book_code = url_match.group(1)
        url_chapter = url_match.group(2)
        
        # Use URL data if we couldn't extract from title
        if not book_name:
            book_name = book_code
        if not chapter_num:
            chapter_num = url_chapter
    else:
        # If we can't parse the URL, use defaults
        book_code = "UNK"
        url_chapter = "0"
    
    # Find all verse elements using the consistent class pattern
    verse_elements = soup.find_all('span', class_=lambda c: c and 'ChapterContent_verse' in c)
    
    if not verse_elements:
        print("Could not find verse elements with class pattern. Trying data-usfm attribute...")
        # Try to find spans with data-usfm attribute
        verse_elements = soup.find_all('span', attrs={'data-usfm': True})
    
    if not verse_elements:
        print("Could not find verse elements with any selector.")
        return None
    
    # Extract verses
    verses = []
    for verse_elem in verse_elements:
        # Get verse number from data-usfm or from label span
        verse_num = ""
        verse_text = ""
        
        # Get verse number from data-usfm attribute (like "GEN.1.5")
        data_usfm = verse_elem.get('data-usfm', '')
        if data_usfm:
            verse_parts = data_usfm.split('.')
            if len(verse_parts) >= 3:
                verse_num = verse_parts[2]
        
        # If no verse number from data-usfm, look for label span
        if not verse_num:
            label_span = verse_elem.find('span', class_=lambda c: c and 'ChapterContent_label' in c)
            if label_span:
                verse_num = label_span.text.strip()
        
        # Get verse text from content span
        content_spans = verse_elem.find_all('span', class_=lambda c: c and 'ChapterContent_content' in c)
        if content_spans:
            # Combine all content spans (sometimes there are multiple)
            verse_text = " ".join(span.text.strip() for span in content_spans if span.text.strip())
        
        # If no specific content span found, use the verse element's text minus the label text
        if not verse_text and verse_elem.text:
            verse_text = verse_elem.text
            # Remove the verse number if it's at the beginning
            if verse_num and verse_text.startswith(verse_num):
                verse_text = verse_text[len(verse_num):].strip()
            # Also remove any nested span texts that might be labels
            for nested_span in verse_elem.find_all('span', class_=lambda c: c and 'ChapterContent_label' in c):
                if nested_span.text in verse_text:
                    verse_text = verse_text.replace(nested_span.text, '').strip()
        
        # Add non-empty verses to our collection
        if verse_num and verse_text:
            verses.append((verse_num, verse_text))
    
    # Get chapter headings/subheadings
    headings = []
    heading_spans = soup.find_all('span', class_=lambda c: c and 'ChapterContent_heading' in c)
    for heading_span in heading_spans:
        heading_text = heading_span.text.strip()
        if heading_text:
            headings.append(heading_text)
    
    return {
        "title": chapter_title or f"Chapter {chapter_num}",
        "book_name": book_name,
        "book_code": book_code,
        "chapter_num": chapter_num,
        "headings": headings,
        "verses": verses
    }

def create_and_setup_docx():
    """
    Create and initialize a new DOCX document
    """
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
    return doc

def add_chapter_to_docx(doc, bible_data, new_book=False):
    """
    Add a chapter to an existing DOCX document
    """
    # Add page break if not the first chapter (or if it's a new book)
    if doc.paragraphs and (len(doc.paragraphs) > 1 or new_book):
        doc.add_page_break()
    
    # Add book title as a major heading if this is a new book
    if new_book and bible_data.get("book_name"):
        book_heading = doc.add_heading(bible_data["book_name"], level=0)
        book_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space after book title
    
    # Add chapter title
    chapter_heading = doc.add_heading(bible_data["title"], level=1)
    chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a separator
    separator = doc.add_paragraph()
    separator.add_run("_" * 40)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add chapter headings/subheadings if any
    for heading in bible_data.get("headings", []):
        heading_para = doc.add_heading(heading, level=2)
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add verses
    for verse_num, verse_text in bible_data["verses"]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        
        # Add verse number in superscript
        verse_num_run = paragraph.add_run(f"{verse_num}")
        verse_num_run.font.bold = True
        verse_num_run.font.size = Pt(9)
        verse_num_run.font.superscript = True
        
        # Add verse text
        verse_text_run = paragraph.add_run(f" {verse_text}")
        verse_text_run.font.size = Pt(11)

def get_next_chapter_url(current_url):
    """
    Determine the URL for the next chapter or book based on known chapter counts
    """
    # Extract book code and chapter from URL
    url_match = re.search(r'/bible/(\d+)/([A-Z0-9]+)\.(\d+)\.(\w+)', current_url)
    if not url_match:
        print("Could not parse URL structure")
        return None
    
    bible_id = url_match.group(1)  # e.g., 2770
    book_code = url_match.group(2)  # e.g., GEN
    chapter_num = int(url_match.group(3))  # e.g., 1
    translation = url_match.group(4)  # e.g., KKB
    
    # Check if we have this book in our dictionary
    if book_code not in CHAPTER_COUNTS:
        print(f"Unknown book code: {book_code}")
        return None
    
    max_chapters = CHAPTER_COUNTS[book_code]
    
    # If we haven't reached the last chapter of the current book
    if chapter_num < max_chapters:
        # Move to the next chapter in the same book
        next_chapter = chapter_num + 1
        next_url = f"https://www.bible.com/bible/{bible_id}/{book_code}.{next_chapter}.{translation}"
        return next_url
    
    # If we've reached the last chapter, move to the next book
    book_index = -1
    for i, (code, _, _) in enumerate(BIBLE_BOOKS):
        if code == book_code:
            book_index = i
            break
    
    if book_index >= 0 and book_index < len(BIBLE_BOOKS) - 1:
        next_book_code = BIBLE_BOOKS[book_index + 1][0]
        next_url = f"https://www.bible.com/bible/{bible_id}/{next_book_code}.1.{translation}"
        print(f"Moving to next book: {next_book_code}")
        return next_url
    else:
        print("Reached the end of the Bible!")
        return None

def main():
    # Create output directory
    output_dir = "bible_text"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate a filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(output_dir, f"Complete_Bible_{timestamp}.docx")
    
    # Create the document
    document = create_and_setup_docx()
    
    # Ask if user wants to start from a specific book and chapter
    custom_start = input("Do you want to start from a specific book and chapter? (y/n): ").lower() == 'y'
    
    # Starting URL (Genesis 1 in KKB Kikamba Bible)
    current_url = "https://www.bible.com/bible/2770/GEN.1.KKB"
    
    if custom_start:
        print("\nAvailable books:")
        for i, (code, name, count) in enumerate(BIBLE_BOOKS):
            print(f"{code}: {name} ({count} chapters)")
            # Print 3 books per line
            if (i + 1) % 3 == 0:
                print()
        
        book_code = input("\nEnter book code (e.g., GEN, LEV): ").upper()
        chapter_num = input("Enter chapter number: ")
        
        # Validate book code
        valid_book = False
        max_chapter = 0
        for code, _, count in BIBLE_BOOKS:
            if code == book_code:
                valid_book = True
                max_chapter = count
                break
        
        if valid_book and chapter_num.isdigit():
            chapter_int = int(chapter_num)
            if 1 <= chapter_int <= max_chapter:
                current_url = f"https://www.bible.com/bible/2770/{book_code}.{chapter_num}.KKB"
                print(f"Starting from {book_code} chapter {chapter_num}")
            else:
                print(f"Invalid chapter number. {book_code} has {max_chapter} chapters. Starting from chapter 1.")
                current_url = f"https://www.bible.com/bible/2770/{book_code}.1.KKB"
        else:
            print("Invalid book code or chapter. Starting from Genesis 1.")
    
    # Flags and counters
    continue_scraping = True
    chapters_scraped = 0
    current_book_code = ""
    
    print("\nStarting to scrape the Bible into a single DOCX document...\n")
    
    try:
        # Continue scraping until we reach the end or user interrupts
        while continue_scraping and current_url:
            try:
                # Scrape current chapter
                bible_data = scrape_bible_text(current_url)
                
                if not bible_data or not bible_data.get("verses"):
                    print(f"Failed to extract verses from {current_url}.")
                    
                    # Try to get the next URL
                    current_url = get_next_chapter_url(current_url)
                    if current_url:
                        print(f"Attempting to continue with: {current_url}")
                        continue
                    else:
                        print("Cannot determine next chapter. Stopping.")
                        break
                
                # Determine if this is a new book
                new_book = bible_data.get("book_code") != current_book_code
                if new_book:
                    current_book_code = bible_data.get("book_code", "")
                    print(f"\nStarting new book: {bible_data.get('book_name')} ({current_book_code})")
                
                # Add chapter to document
                add_chapter_to_docx(document, bible_data, new_book)
                
                chapters_scraped += 1
                print(f"Added chapter: {bible_data['title']} (Total: {chapters_scraped})")
                
                # Save progress periodically (every 10 chapters)
                if chapters_scraped % 10 == 0:
                    print(f"Saving progress... (Current chapter count: {chapters_scraped})")
                    document.save(output_file)
                
                # Get next URL
                current_url = get_next_chapter_url(current_url)
                if not current_url:
                    print("Reached the end of the Bible.")
                    continue_scraping = False
                
                # Respectful delay between requests
                if continue_scraping:
                    print("Waiting before next chapter...")
                    time.sleep(3)
                    
            except KeyboardInterrupt:
                raise  # Re-raise to be caught by outer try/except
            except Exception as e:
                print(f"Error: {e}")
                print("Trying to continue to next chapter...")
                
                # Try to continue with the next URL
                current_url = get_next_chapter_url(current_url)
                if current_url:
                    print(f"Attempting to continue with: {current_url}")
                else:
                    print("Cannot determine next chapter. Stopping.")
                    break
        
        # Save the final document
        print(f"\nSaving complete document with {chapters_scraped} chapters...")
        document.save(output_file)
            
        print(f"\nDone! Bible with {chapters_scraped} chapters saved to:")
        print(f"{output_file}")
        
    except KeyboardInterrupt:
        print("\nScraping interrupted by user.")
        # Save what we have so far
        print(f"Saving document with {chapters_scraped} chapters...")
        document.save(output_file)
        print(f"Partial Bible with {chapters_scraped} chapters saved to:")
        print(f"{output_file}")

if __name__ == "__main__":
    main()